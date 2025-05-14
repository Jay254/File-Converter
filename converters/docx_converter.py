import os
from docx import Document
from docx2pdf import convert
from PIL import Image
import comtypes.client
import pythoncom

def convert_docx(filepath, target_format):
    """
    Convert DOCX to various formats
    """
    output_path = os.path.splitext(filepath)[0] + f'.{target_format}'
    
    if target_format == 'pdf':
        # DOCX to PDF conversion
        convert(filepath, output_path)
        
    elif target_format == 'txt':
        # DOCX to TXT conversion
        doc = Document(filepath)
        with open(output_path, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                f.write(para.text + '\n')
                
    elif target_format == 'html':
        # DOCX to HTML conversion
        doc = Document(filepath)
        html_content = ['<html><body>']
        for para in doc.paragraphs:
            html_content.append(f'<p>{para.text}</p>')
        html_content.append('</body></html>')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_content))
            
    else:
        raise ValueError(f"Conversion from DOCX to {target_format} is not supported")
        
    return output_path 